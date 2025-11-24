import io
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aiogram import Bot
from aiogram.types import FSInputFile

from src.data.repositories.specialist_repository import specialist_crud


async def download_qr_image(bot: Bot, file_id: str) -> bytes:
    """
    Скачивает QR-код по file_id из Telegram.
    
    Args:
        bot: Экземпляр бота
        file_id: ID файла в Telegram
        
    Returns:
        Байты изображения
    """
    file_info = await bot.get_file(file_id)
    bio = io.BytesIO()
    await bot.download_file(file_info.file_path, destination=bio)
    bio.seek(0)
    return bio.read()


async def generate_specialists_word(bot: Bot, organization: str = None) -> str:
    """
    Генерирует Word документ со списком специалистов.
    
    Args:
        bot: Экземпляр бота
        organization: Фильтр по организации (необязательно)
        
    Returns:
        Путь к созданному файлу
    """
    # Получаем список специалистов
    if organization:
        specialists = await specialist_crud.get_list(organization=organization)
        filename = f"specialists_{organization.replace(' ', '_')}.docx"
    else:
        specialists = await specialist_crud.get_list()
        filename = "specialists_all.docx"
    
    # Создаем документ
    doc = Document()
    
    # Добавляем заголовок
    title = doc.add_heading('Список специалистов', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if organization:
        subtitle = doc.add_heading(f'Организация: {organization}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Добавляем каждого специалиста
    for idx, spec in enumerate(specialists, 1):
        # Разделитель между специалистами
        if idx > 1:
            doc.add_page_break()
        
        # Номер и ФИО
        heading = doc.add_heading(f'{idx}. {spec.fullname}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о специалисте
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        # Должность
        row = info_table.add_row()
        row.cells[0].text = '📋 Должность:'
        row.cells[1].text = spec.position or 'Не указано'
        
        # Организация
        row = info_table.add_row()
        row.cells[0].text = '🏢 Организация:'
        row.cells[1].text = spec.organization
        
        # Отдел
        if spec.department and spec.department != '-':
            row = info_table.add_row()
            row.cells[0].text = '📁 Отдел:'
            row.cells[1].text = spec.department
        
        # Ссылка
        if spec.link:
            row = info_table.add_row()
            row.cells[0].text = '🔗 Ссылка:'
            row.cells[1].text = spec.link
        
        doc.add_paragraph()
        
        # QR-код
        if spec.qr:
            try:
                qr_heading = doc.add_paragraph('QR-код для связи:')
                qr_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                qr_heading.runs[0].bold = True
                
                # Скачиваем и добавляем QR-код
                qr_bytes = await download_qr_image(bot, spec.qr)
                
                # Сохраняем временно
                temp_qr_path = f'temp_qr_{spec.id}.png'
                with open(temp_qr_path, 'wb') as f:
                    f.write(qr_bytes)
                
                # Добавляем изображение в документ
                qr_paragraph = doc.add_paragraph()
                qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = qr_paragraph.add_run()
                run.add_picture(temp_qr_path, width=Inches(3))
                
                # Удаляем временный файл
                os.remove(temp_qr_path)
                
            except Exception as e:
                doc.add_paragraph(f'⚠️ QR-код недоступен')
                print(f'Ошибка загрузки QR для {spec.fullname}: {e}')
        else:
            doc.add_paragraph('⚠️ QR-код не сгенерирован')
    
    # Сохраняем документ
    os.makedirs('src/files/exports', exist_ok=True)
    filepath = f'src/files/exports/{filename}'
    doc.save(filepath)
    
    return filepath


async def export_specialists_to_word(bot: Bot, chat_id: int, organization: str = None):
    """
    Экспортирует специалистов в Word и отправляет файл пользователю.
    
    Args:
        bot: Экземпляр бота
        chat_id: ID чата для отправки файла
        organization: Фильтр по организации (необязательно)
    """
    try:
        # Отправляем сообщение о начале генерации
        status_msg = await bot.send_message(
            chat_id=chat_id,
            text="⏳ <b>Генерация Word документа...</b>\n\nПожалуйста, подождите.",
            parse_mode="HTML"
        )
        
        # Генерируем документ
        filepath = await generate_specialists_word(bot, organization)
        
        # Отправляем файл
        document = FSInputFile(filepath)
        await bot.send_document(
            chat_id=chat_id,
            document=document,
            caption="✅ <b>Список специалистов</b>\n\nДокумент содержит информацию о специалистах с QR-кодами.",
            parse_mode="HTML"
        )
        
        # Удаляем сообщение о статусе
        await status_msg.delete()
        
        # Удаляем файл после отправки
        os.remove(filepath)
        
    except Exception as e:
        await bot.send_message(
            chat_id=chat_id,
            text=f"❌ <b>Ошибка при генерации документа:</b>\n\n<code>{str(e)}</code>",
            parse_mode="HTML"
        )
        print(f'Ошибка экспорта в Word: {e}')
